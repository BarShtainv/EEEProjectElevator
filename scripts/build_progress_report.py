#!/usr/bin/env python3
"""Build a polished progress-report DOCX from a Markdown source."""

from __future__ import annotations

import argparse
from pathlib import Path
import re
import subprocess
import tempfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_SOURCE = ROOT / "report" / "progress" / "advancement_report_1.md"
DEFAULT_OUTPUT = ROOT / "report" / "progress" / "advancement_report_1.docx"

PAGE_BREAK = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""


def _set_page_number_start(section, start: int) -> None:
    page_number = section._sectPr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        section._sectPr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    for kind, value in (
        ("begin", None),
        (None, " PAGE "),
        ("separate", None),
        ("end", None),
    ):
        if kind is not None:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), kind)
        else:
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = value
        run._r.append(node)


def _repeat_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    row_properties.append(header)


def _prevent_row_split(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))


def _shade_cell(cell, fill: str) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _configure_reference_doc(path: Path, report_title: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    section.different_first_page_header_footer = True
    _set_page_number_start(section, 0)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.12
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.widow_control = True

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(26)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(82)
    title.paragraph_format.space_after = Pt(34)

    for name, size, colour, before, after in (
        ("Heading 1", 16, RGBColor(31, 78, 121), 14, 6),
        ("Heading 2", 12.5, RGBColor(47, 84, 150), 10, 4),
        ("Heading 3", 11, RGBColor(68, 68, 68), 8, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = colour
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = doc.styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(8.8)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(60, 60, 60)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_together = True

    if "Compact" not in doc.styles:
        compact = doc.styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        compact = doc.styles["Compact"]
    compact.font.name = "Aptos"
    compact.font.size = Pt(8.2)
    compact.paragraph_format.space_before = Pt(1.2)
    compact.paragraph_format.space_after = Pt(1.2)
    compact.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = f"{report_title}  |  Final Project Controlled Floor Elevator"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(7.5)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    _add_page_field(section.footer.paragraphs[0])
    doc.save(path)


def _production_markdown(source: str, report_title: str) -> str:
    lines = source.splitlines()
    if not lines or lines[0].strip() != f"# {report_title}":
        raise ValueError(f"source must begin with '# {report_title}'")
    body = "\n".join(lines[1:]).lstrip()
    marker = "## 1. Introduction and Project Objective"
    if marker not in body:
        raise ValueError("expected first numbered report section is missing")
    cover, report_body = body.split(marker, 1)
    return f"""---
title: "{report_title}"
lang: en-GB
---

{cover.rstrip()}

{PAGE_BREAK}

## 1. Introduction and Project Objective
{report_body.lstrip()}
"""


def _postprocess_docx(path: Path, report_title: str) -> None:
    doc = Document(path)
    doc.core_properties.title = f"{report_title} — Final Project Controlled Floor Elevator"
    doc.core_properties.author = "Bar Shtainvortzel"
    doc.core_properties.subject = "B.Sc. Electrical and Electronics Engineering advancement report"
    doc.core_properties.keywords = "access control, elevator, authorization, software design, verification planning"

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.9)
        section.left_margin = Cm(2.1)
        section.right_margin = Cm(2.1)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.9)
        section.different_first_page_header_footer = True
        _set_page_number_start(section, 0)

    cover_complete = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif not cover_complete and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
        if text == "1. Introduction and Project Objective":
            cover_complete = True
        if re.match(r"^Table \d+\.", text):
            paragraph.style = doc.styles["Caption"]
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        try:
            table.style = "Light Shading Accent 1"
        except KeyError:
            table.style = "Table Grid"
        for row_index, row in enumerate(table.rows):
            _prevent_row_split(row)
            if row_index == 0:
                _repeat_header(row)
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.save(path)


def build(source: Path, output: Path, report_title: str) -> None:
    text = source.read_text(encoding="utf-8")
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="eee-progress-report-") as temporary_name:
        temporary = Path(temporary_name)
        reference = temporary / "reference.docx"
        production = temporary / "report.md"
        _configure_reference_doc(reference, report_title)
        production.write_text(_production_markdown(text, report_title), encoding="utf-8")
        subprocess.run(
            [
                "pandoc",
                str(production),
                "--from=markdown+raw_attribute",
                "--to=docx",
                f"--reference-doc={reference}",
                "--resource-path",
                str(source.parent) + ":" + str(ROOT),
                "--metadata",
                f"pagetitle={report_title}",
                "--output",
                str(output),
            ],
            cwd=ROOT,
            check=True,
        )
    _postprocess_docx(output, report_title)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=DEFAULT_SOURCE)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--title", default="ADVANCEMENT REPORT 1")
    args = parser.parse_args()
    build(args.source.resolve(), args.output.resolve(), args.title)


if __name__ == "__main__":
    main()
