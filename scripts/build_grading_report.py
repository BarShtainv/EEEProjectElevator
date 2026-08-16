#!/usr/bin/env python3
"""Build the grading-review DOCX from the canonical Markdown report."""

from __future__ import annotations

import argparse
from pathlib import Path
import re
import subprocess
import tempfile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "report" / "final_report.md"
OUTPUT = ROOT / "report" / "final_report_grading_draft.docx"

PAGE_BREAK = """```{=openxml}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```"""

TOC_FIELD = """```{=openxml}
<w:p><w:r><w:fldChar w:fldCharType="begin" w:dirty="true"/></w:r><w:r><w:instrText xml:space="preserve"> TOC \\o "1-2" \\h \\z \\u </w:instrText></w:r><w:r><w:fldChar w:fldCharType="separate"/></w:r><w:r><w:t>Right-click and select Update Field if the table of contents is not populated.</w:t></w:r><w:r><w:fldChar w:fldCharType="end"/></w:r></w:p>
```"""


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_row_no_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _set_page_number_start(section, start: int) -> None:
    sect_pr = section._sectPr
    page_number = sect_pr.find(qn("w:pgNumType"))
    if page_number is None:
        page_number = OxmlElement("w:pgNumType")
        sect_pr.append(page_number)
    page_number.set(qn("w:start"), str(start))


def _add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, end))


def _configure_reference_doc(path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.25)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(1.1)
    section.footer_distance = Cm(1.1)
    section.different_first_page_header_footer = True
    _set_page_number_start(section, 0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(25)
    title.font.bold = True
    title.font.color.rgb = RGBColor(31, 78, 121)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    title.paragraph_format.space_after = Pt(30)

    heading_specs = {
        "Heading 1": (17, RGBColor(31, 78, 121), Pt(16), Pt(8)),
        "Heading 2": (13.5, RGBColor(47, 84, 150), Pt(13), Pt(5)),
        "Heading 3": (11.5, RGBColor(68, 68, 68), Pt(10), Pt(3)),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = before
        style.paragraph_format.space_after = after
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    caption = styles["Caption"]
    caption.font.name = "Aptos"
    caption.font.size = Pt(9)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(60, 60, 60)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    if "Source Code" not in styles:
        code = styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code.font.size = Pt(8.5)
        code.paragraph_format.left_indent = Cm(0.45)
        code.paragraph_format.right_indent = Cm(0.25)
        code.paragraph_format.space_before = Pt(3)
        code.paragraph_format.space_after = Pt(5)

    if "Compact" not in styles:
        compact = styles.add_style("Compact", WD_STYLE_TYPE.PARAGRAPH)
    else:
        compact = styles["Compact"]
    compact.font.name = "Aptos"
    compact.font.size = Pt(8.2)
    compact.paragraph_format.space_before = Pt(1.5)
    compact.paragraph_format.space_after = Pt(1.5)
    compact.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.text = "Final Project Controlled Floor Elevator"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Aptos"
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    _add_page_field(section.footer.paragraphs[0])
    doc.save(path)


def _production_markdown(source: str) -> str:
    lines = source.splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError("report must begin with a level-one title")
    title = lines[0][2:].strip()
    body = "\n".join(lines[1:]).lstrip()

    abstract_marker = "## 1. Abstract"
    introduction_marker = "## 2. Introduction"
    if abstract_marker not in body or introduction_marker not in body:
        raise ValueError("expected abstract and introduction headings are missing")
    cover, rest = body.split(abstract_marker, 1)
    cover = cover.replace("  \n", "\n\n")
    abstract_body, remaining = rest.split(introduction_marker, 1)

    figures = re.findall(r"!\[Figure (\d+)\.\s*([^\]]+)\]", source)
    tables = re.findall(r"\*\*Table (\d+)\.\s*([^*]+)\*\*", source)
    figure_list = "\n".join(f"{number}. {caption}" for number, caption in figures)
    table_list = "\n".join(f"{number}. {caption.strip()}" for number, caption in tables)

    return f"""---
title: "{title}"
lang: en-GB
---

{cover.rstrip()}

{PAGE_BREAK}

## 1. Abstract
{abstract_body.rstrip()}

{PAGE_BREAK}

**Table of Contents**

{TOC_FIELD}

{PAGE_BREAK}

**List of Figures**

{figure_list}

**List of Tables**

{table_list}

{PAGE_BREAK}

## 2. Introduction
{remaining.lstrip()}
"""


def _postprocess_docx(path: Path) -> None:
    doc = Document(path)
    doc.core_properties.title = "Final Project Controlled Floor Elevator"
    doc.core_properties.author = "Bar Shtainvortzel"
    doc.core_properties.subject = "B.Sc. Electrical and Electronics Engineering final project report"
    doc.core_properties.keywords = "access control, RFID, elevator, authorization, simulation, verification"

    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.25)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.35)
        section.right_margin = Cm(2.35)
        section.header_distance = Cm(1.1)
        section.footer_distance = Cm(1.1)
        section.different_first_page_header_footer = True
        _set_page_number_start(section, 0)

    cover_complete = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif not cover_complete and text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
        if text == "1. Abstract":
            cover_complete = True
        if re.match(r"^(Table of Contents|List of Figures|List of Tables)$", text):
            paragraph.style = doc.styles["Heading 1"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if re.match(r"^Table \d+\.", text):
            paragraph.style = doc.styles["Caption"]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Caption":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = False
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.keep_together = True
        if paragraph._p.xpath(".//w:drawing"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_with_next = True

    doc.save(path)


def build(output: Path) -> None:
    source = SOURCE.read_text(encoding="utf-8")
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix="eee-report-build-") as tmp_name:
        tmp = Path(tmp_name)
        reference = tmp / "reference.docx"
        production_md = tmp / "report.md"
        _configure_reference_doc(reference)
        production_md.write_text(_production_markdown(source), encoding="utf-8")
        subprocess.run(
            [
                "pandoc",
                str(production_md),
                "--from=markdown+raw_attribute",
                "--to=docx",
                f"--reference-doc={reference}",
                "--resource-path",
                str(ROOT / "report") + ":" + str(ROOT),
                "--metadata",
                "pagetitle=Final Project Controlled Floor Elevator",
                "--output",
                str(output),
            ],
            cwd=ROOT,
            check=True,
        )
    _postprocess_docx(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--output", type=Path, default=OUTPUT)
    args = parser.parse_args()
    build(args.output.resolve())


if __name__ == "__main__":
    main()
